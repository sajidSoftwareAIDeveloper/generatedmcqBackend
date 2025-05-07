
# :: 1. Create virtual environment
# py -3.10 -m venv spacy_env

# :: 2. Activate it
# spacy_env\Scripts\activate

# python -m pip install --upgrade pip setuptools wheel
# python -m pip install spacy transformers pytesseract opencv-python pdfplumber python-docx django djangorestframework
# python -m spacy download en_core_web_sm
# pip install torch
# pip install tiktoken protobuf
# pip install djangorestframework
# pip install django

# sudo apt-get install tesseract-ocr-hin      for hindi hin



from PIL import Image
import pytesseract
import cv2
import pdfplumber
from docx import Document

# Set the path to tesseract.exe (custom install path)
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\sajid.anwar\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

def image_to_text(file):
    # print('welcome to file')
    image = Image.open(file)
    text = pytesseract.image_to_string(image ,lang='eng')
    return text


# with layout 
def pdf_to_text(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


# def docx_to_text(file):
#     try:
#         doc = Document(file)
#         return "\n".join([para.text for para in doc.paragraphs])
#     except Exception as e:
#         return f"Error reading DOCX file: {str(e)}"

def docx_to_text(file):
    try:
        file.seek(0)  # Reset pointer if needed
        doc = Document(file)  # ✅ Correct usage
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

def uploadFile(file):
    file_type = file.content_type
    # print('>>>>',file_type,file)
    
    if "pdf" in file_type:
        # print('pdf')
        return pdf_to_text(file)
    elif "image" in file_type:
        # print('image')
        return image_to_text(file)
    elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in file_type:
        return docx_to_text(file)

    else:
        return "Unsupported file type"